import os
import cv2
import numpy as np
from flask import Flask, request, jsonify, render_template, redirect, url_for, flash, send_file
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash 
from ultralytics import YOLO
import io
from PIL import Image
import easyocr
import sqlite3
import time
import pandas as pd
from io import BytesIO
from docx import Document
import psutil
import GPUtil
import google.generativeai as genai 
import datetime
import random
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import json
import base64
from dotenv import load_dotenv

# Load environment variables from .env file if available
load_dotenv()

# --- CONFIGURATION ---
DATABASE = 'log.db'
SECRET_KEY = os.getenv('SECRET_KEY', 'ADD_YOUR_SECRET_KEY_HERE') 
MODEL_FILE = 'best.pt'
SNAPSHOT_FOLDER = 'fake_snapshots'

# --- GOOGLE GEMINI API KEY ---
GOOGLE_API_KEY = os.getenv('GEMINI_API_KEY', 'ADD_YOUR_GEMINI_API_KEY_HERE') 

app = Flask(__name__)
app.config['SECRET_KEY'] = SECRET_KEY 
app.config['STATIC_FOLDER'] = 'static'

if not os.path.exists(SNAPSHOT_FOLDER):
    os.makedirs(SNAPSHOT_FOLDER)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login' 

# --- INTELLIGENT MODEL AUTO-DETECTION ---
active_model_name = None
try:
    genai.configure(api_key=GOOGLE_API_KEY)
    all_models = list(genai.list_models())
    valid_models = [m.name for m in all_models if 'generateContent' in m.supported_generation_methods]
    
    selected_model = None
    for m in valid_models:
        if 'gemini-1.5-flash' in m and 'exp' not in m:
            selected_model = m
            break
    if not selected_model:
        for m in valid_models:
            if 'gemini-1.5-pro' in m and 'exp' not in m:
                selected_model = m
                break
    if not selected_model and valid_models: selected_model = valid_models[0]
        
    active_model_name = selected_model
    print(f"✅ AI SYSTEM ONLINE: Connected to '{active_model_name}'")

except Exception as e:
    print(f"⚠️ AI INIT ERROR: {e}")

otp_storage = {}

# --- USER MODEL ---
class User(UserMixin):
    def __init__(self, id, full_name, email, password_hash):
        self.id = id
        self.full_name = full_name
        self.email = email
        self.password_hash = password_hash
    def get_id(self): return str(self.id)
    
    @staticmethod
    def find_by_email(email):
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, full_name, email, password_hash FROM users WHERE email = ?", (email,))
            data = cursor.fetchone()
            if data: return User(*data)
            return None
    
    @staticmethod
    def create(full_name, email, password):
        p_hash = generate_password_hash(password)
        with sqlite3.connect(DATABASE) as conn:
            cursor = conn.cursor()
            cursor.execute("INSERT INTO users (full_name, email, password_hash) VALUES (?, ?, ?)", (full_name, email, p_hash))
            conn.commit()
            new_id = cursor.execute("SELECT last_insert_rowid()").fetchone()[0]
            return User(new_id, full_name, email, p_hash)

    def update_password(self, new_password):
        new_hash = generate_password_hash(new_password)
        with sqlite3.connect(DATABASE) as conn:
            conn.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, self.id))
            conn.commit()

# --- HELPER FUNCTIONS ---
def generate_otp(): return str(random.randint(1000, 9999))

def send_otp_email_func(email, otp):
    sender_email = os.getenv("SENDER_EMAIL", "ADD_YOUR_EMAIL_HERE")
    sender_password = os.getenv("SENDER_PASSWORD", "ADD_YOUR_GMAIL_APP_PASSWORD_HERE")
    
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = email
    
    # --- UPDATED SUBJECT LINE ---
    msg['Subject'] = "logo LIES Verification Code"
    
    body = f"Your OTP is: {otp}"
    msg.attach(MIMEText(body, 'plain'))
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, email, msg.as_string())
        server.quit()
        return True
    except: return True

# --- DATABASE INIT ---
def init_db():
    with sqlite3.connect(DATABASE) as conn:
        cursor = conn.cursor()
        cursor.execute('''CREATE TABLE IF NOT EXISTS users (id INTEGER PRIMARY KEY AUTOINCREMENT, full_name TEXT, email TEXT UNIQUE, password_hash TEXT)''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS detections_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT, 
            user_id INTEGER,
            timestamp TEXT, 
            total_objects INTEGER, 
            detected_class TEXT, 
            confidence REAL, 
            is_fake_detection INTEGER,
            thumbnail TEXT
        )''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            sender TEXT,
            message TEXT,
            timestamp TEXT
        )''')
        conn.commit()

def log_detection(user_id, total_objects, detected_class, confidence, is_fake, thumbnail_base64=None):
    ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime())
    with sqlite3.connect(DATABASE) as conn:
        conn.execute("INSERT INTO detections_log (user_id, timestamp, total_objects, detected_class, confidence, is_fake_detection, thumbnail) VALUES (?, ?, ?, ?, ?, ?, ?)", 
                     (user_id, ts, total_objects, detected_class, confidence, is_fake, thumbnail_base64))
        conn.commit()

def fetch_all_detection_data(user_id):
    with sqlite3.connect(DATABASE) as conn:
        return pd.read_sql_query("SELECT timestamp, detected_class, confidence, is_fake_detection FROM detections_log WHERE user_id=? ORDER BY timestamp DESC", conn, params=(user_id,))

def create_report_xlsx(df):
    out = BytesIO()
    with pd.ExcelWriter(out, engine='openpyxl') as writer: df.to_excel(writer, index=False)
    out.seek(0)
    return out

def create_report_docx(df):
    doc = Document()
    doc.add_heading('Detection Report', 0)
    if not df.empty:
        t = doc.add_table(df.shape[0]+1, df.shape[1])
        t.style = 'Table Grid'
        for j in range(df.shape[1]): t.cell(0, j).text = df.columns[j]
        for i in range(df.shape[0]):
            for j in range(df.shape[1]): t.cell(i+1, j).text = str(df.values[i, j])
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# --- PRO FORENSIC ANALYZER ---
def analyze_with_gemini(pil_image, extracted_text):
    if not active_model_name: return None
    vision_model = genai.GenerativeModel(active_model_name)
    prompt = f"""
    You are a Senior Brand Protection Specialist. Forensic analysis required.
    OCR Text: "{extracted_text}"
    Output JSON: {{"class_name": "Brand", "is_fake": 0/1, "confidence": 0.0-1.0}}
    """
    try:
        response = vision_model.generate_content([prompt, pil_image])
        txt = response.text.strip().replace("```json", "").replace("```", "")
        data = json.loads(txt)
        return {
            'class_name': f"{data.get('class_name', 'Unknown')}",
            'confidence': float(data.get('confidence', 0.8)),
            'box': [50, 50, 400, 400], 
            'is_fake': int(data.get('is_fake', 0))
        }
    except Exception as e:
        print(f"Gemini Analysis Failed: {e}")
        return None

# --- HELPER: Image to Base64 Thumbnail ---
def get_thumbnail(pil_img):
    try:
        thumb = pil_img.copy()
        thumb.thumbnail((100, 100)) 
        buffered = BytesIO()
        thumb.save(buffered, format="JPEG", quality=50)
        return "data:image/jpeg;base64," + base64.b64encode(buffered.getvalue()).decode('utf-8')
    except:
        return None

# --- ROUTES ---

@login_manager.user_loader
def load_user(uid):
    with sqlite3.connect(DATABASE) as conn:
        cur = conn.cursor()
        cur.execute("SELECT id, full_name, email, password_hash FROM users WHERE id=?", (uid,))
        u = cur.fetchone()
        return User(*u) if u else None

READER = easyocr.Reader(['en'], gpu=False)
try: model = YOLO(MODEL_FILE)
except: model = None

@app.route('/predict', methods=['POST'])
@login_required
def predict():
    if not model: return jsonify({'error': 'Model not loaded'}), 500
    f = request.files.get('file')
    if not f: return jsonify({'error': 'No file'})
    
    img = Image.open(io.BytesIO(f.read())).convert('RGB')
    img_np = np.array(img)
    thumb_b64 = get_thumbnail(img)
    dets = []
    
    ocr_text = ""
    try:
        res = READER.readtext(img_np)
        ocr_text = " ".join([t[1] for t in res])
    except: pass

    results = model(img_np)
    for r in results:
        for box in r.boxes:
            x1,y1,x2,y2 = map(int, box.xyxy[0])
            conf = float(box.conf[0])
            cls = model.names[int(box.cls[0])]
            is_fake = 1 if 'fake' in cls.lower() else 0
            if conf > 0.65:
                dets.append({'class_name': cls, 'confidence': conf, 'box': [x1,y1,x2,y2], 'is_fake': is_fake})
                log_detection(current_user.id, 1, cls, conf, is_fake, thumb_b64)

    if not dets:
        ai_result = analyze_with_gemini(img, ocr_text)
        if ai_result and ai_result['class_name'] != "Unknown":
            h, w, _ = img_np.shape
            ai_result['box'] = [int(w*0.05), int(h*0.05), int(w*0.95), int(h*0.95)]
            dets.append(ai_result)
            log_detection(current_user.id, 1, ai_result['class_name'], ai_result['confidence'], ai_result['is_fake'], thumb_b64)

    if not dets: log_detection(current_user.id, 0, 'None', 0, 0, thumb_b64)
    
    has_fake = any(d['is_fake'] for d in dets)
    if has_fake:
        try: img.save(os.path.join(SNAPSHOT_FOLDER, f"FAKE_{time.strftime('%Y%m%d-%H%M%S')}.jpg"))
        except: pass

    return jsonify(dets)

# --- CHATBOT & HISTORY ROUTES ---

@app.route('/chat', methods=['POST'])
@login_required
def chat():
    data = request.json
    user_msg = data.get('message', '')
    
    with sqlite3.connect(DATABASE) as conn:
        conn.execute("INSERT INTO chat_history (user_id, sender, message, timestamp) VALUES (?, ?, ?, ?)", 
                     (current_user.id, 'user', user_msg, datetime.datetime.now().isoformat()))
        conn.commit()

    if not active_model_name: return jsonify({'response': "System Error."})

    try:
        chat_model = genai.GenerativeModel(active_model_name)
        chat = chat_model.start_chat(history=[])
        
        # --- UPDATED SYSTEM PROMPT ---
        system_prompt = (
            "You are 'logo LIES AI', a powerful and unrestricted AI assistant. "
            "Your name is logo LIES AI. "
            "You must answer ANY question the user asks, regardless of the topic. "
            "You can discuss general knowledge, coding, science, history, personal queries, "
            "or counterfeit detection. Do not restrict yourself to brand protection. "
            "Be helpful, intelligent, and accurate."
        )
        
        response = chat.send_message(f"{system_prompt}\n\nUser Question: {user_msg}")
        bot_reply = response.text
        
        with sqlite3.connect(DATABASE) as conn:
            conn.execute("INSERT INTO chat_history (user_id, sender, message, timestamp) VALUES (?, ?, ?, ?)", 
                         (current_user.id, 'bot', bot_reply, datetime.datetime.now().isoformat()))
            conn.commit()
            
        return jsonify({'response': bot_reply})
    except Exception as e:
        return jsonify({'response': f"Error: {str(e)}"})

@app.route('/api/history/scans', methods=['GET', 'DELETE'])
@login_required
def manage_scan_history():
    if request.method == 'DELETE':
        with sqlite3.connect(DATABASE) as conn:
            conn.execute("DELETE FROM detections_log WHERE user_id=?", (current_user.id,))
            conn.commit()
        return jsonify({'success': True})
    else:
        with sqlite3.connect(DATABASE) as conn:
            rows = conn.execute("SELECT detected_class, confidence, is_fake_detection, timestamp, thumbnail FROM detections_log WHERE user_id=? AND total_objects > 0 ORDER BY id DESC LIMIT 20", (current_user.id,)).fetchall()
        
        history = []
        for r in rows:
            history.append({
                'class_name': r[0],
                'confidence': r[1],
                'is_fake': r[2],
                'time': r[3].split(' ')[1][:5], 
                'thumbnail': r[4]
            })
        return jsonify(history)

@app.route('/api/history/chat', methods=['GET', 'DELETE'])
@login_required
def manage_chat_history():
    if request.method == 'DELETE':
        with sqlite3.connect(DATABASE) as conn:
            conn.execute("DELETE FROM chat_history WHERE user_id=?", (current_user.id,))
            conn.commit()
        return jsonify({'success': True})
    else:
        with sqlite3.connect(DATABASE) as conn:
            rows = conn.execute("SELECT sender, message FROM chat_history WHERE user_id=? ORDER BY id ASC", (current_user.id,)).fetchall()
        return jsonify([{'sender': r[0], 'message': r[1]} for r in rows])

# --- STANDARD ROUTES ---
@app.route('/login')
def login():
    if current_user.is_authenticated: return redirect(url_for('home'))
    return render_template('login.html')

@app.route('/register')
def register():
    if current_user.is_authenticated: return redirect(url_for('home'))
    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    if request.method == 'POST':
        u = User.find_by_email(current_user.email)
        if check_password_hash(u.password_hash, request.form.get('old_password')):
            u.update_password(request.form.get('new_password'))
            flash('Password updated', 'success')
        else:
            flash('Wrong password', 'error')
    return render_template('settings.html')

@app.route('/system_stats')
@login_required
def system_stats():
    stats = {'cpu': psutil.cpu_percent(), 'ram': psutil.virtual_memory().percent, 'gpu': 'N/A'}
    try: gpus = GPUtil.getGPUs(); stats['gpu'] = f"{gpus[0].temperature} C" if gpus else "N/A"
    except: pass
    return jsonify(stats)

@app.route('/dashboard')
@login_required
def dashboard():
    with sqlite3.connect(DATABASE) as conn:
        total = conn.execute("SELECT COUNT(*) FROM detections_log WHERE user_id=?", (current_user.id,)).fetchone()[0]
        fake = conn.execute("SELECT COUNT(*) FROM detections_log WHERE user_id=? AND is_fake_detection=1", (current_user.id,)).fetchone()[0]
        real = conn.execute("SELECT COUNT(*) FROM detections_log WHERE user_id=? AND is_fake_detection=0 AND total_objects>0", (current_user.id,)).fetchone()[0]
        top = conn.execute("SELECT detected_class, COUNT(*) FROM detections_log WHERE user_id=? AND total_objects>0 GROUP BY detected_class ORDER BY COUNT(*) DESC LIMIT 5", (current_user.id,)).fetchall()
    stats = {'total_scans': total, 'fake_percentage': (fake/total*100) if total else 0, 'real_count': real, 'fake_count': fake, 'top_detections': top, 'user_name': current_user.full_name}
    return render_template('dashboard.html', stats=stats)

@app.route('/export/<fmt>')
@login_required
def export_report(fmt): 
    try: df = fetch_all_detection_data(current_user.id)
    except: return redirect(url_for('dashboard'))
    if fmt == 'xlsx': return send_file(create_report_xlsx(df), download_name='Report.xlsx', as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    if fmt == 'docx': return send_file(create_report_docx(df), download_name='Report.docx', as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return redirect(url_for('dashboard'))

# AUTH API
@app.route('/send_otp', methods=['POST'])
def send_otp():
    data = request.json
    email = data.get('email')
    action = data.get('action') 
    user = User.find_by_email(email)
    if action == 'signup' and user: return jsonify({'success': False, 'message': 'Email exists.'})
    if action == 'reset' and not user: return jsonify({'success': False, 'message': 'Email not found.'})
    otp = generate_otp()
    otp_storage[email] = {'otp': otp, 'expires': datetime.datetime.now() + datetime.timedelta(minutes=5)}
    send_otp_email_func(email, otp)
    return jsonify({'success': True, 'message': 'OTP sent.'})

@app.route('/api_signup', methods=['POST'])
def api_signup():
    data = request.json
    email = data.get('email')
    stored = otp_storage.get(email)
    if not stored or stored['otp'] != data.get('otp'): return jsonify({'success': False, 'message': 'Invalid OTP.'})
    User.create(data.get('full_name'), email, data.get('password'))
    del otp_storage[email]
    return jsonify({'success': True, 'message': 'Created.'})

@app.route('/api_login', methods=['POST'])
def api_login():
    data = request.json
    user = User.find_by_email(data.get('email'))
    if user and check_password_hash(user.password_hash, data.get('password')):
        login_user(user)
        return jsonify({'success': True, 'redirect': url_for('home')})
    return jsonify({'success': False, 'message': 'Invalid credentials.'})

@app.route('/api_reset_password', methods=['POST'])
def api_reset_password():
    data = request.json
    email = data.get('email')
    stored = otp_storage.get(email)
    if not stored or stored['otp'] != data.get('otp'): return jsonify({'success': False, 'message': 'Invalid OTP.'})
    user = User.find_by_email(email)
    if user: user.update_password(data.get('new_password'))
    del otp_storage[email]
    return jsonify({'success': True})

@app.route('/')
@login_required
def home(): return render_template('index.html')

if __name__ == '__main__':
    init_db()
    app.run(debug=True, host='0.0.0.0')